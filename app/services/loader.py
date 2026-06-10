"""Document Loader - Supports PDF, DOCX, TXT, MD"""
import os
from pathlib import Path
from typing import List, Optional
import logging

logger = logging.getLogger(__name__)

class Document:
    """Represents a loaded document"""
    def __init__(self, content: str, metadata: dict):
        self.page_content = content
        self.metadata = metadata

class DocumentLoader:
    """Load various document formats"""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html"}
    
    @staticmethod
    def load_file(file_path: str) -> List[Document]:
        """Load a single file and return Document objects"""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext not in DocumentLoader.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")
        
        if ext == ".pdf":
            return DocumentLoader._load_pdf(file_path)
        elif ext == ".docx":
            return DocumentLoader._load_docx(file_path)
        elif ext in (".txt", ".md", ".html"):
            return DocumentLoader._load_text(file_path)
    
    @staticmethod
    def _load_pdf(file_path: str) -> List[Document]:
        """Load PDF using pypdf"""
        from pypdf import PdfReader
        
        documents = []
        reader = PdfReader(file_path)
        
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                documents.append(Document(
                    content=text,
                    metadata={
                        "source": file_path,
                        "page": i + 1,
                        "total_pages": len(reader.pages),
                        "document_type": "pdf"
                    }
                ))
        
        logger.info(f"Loaded PDF: {file_path}, {len(documents)} pages")
        return documents
    
    @staticmethod
    def _load_docx(file_path: str) -> List[Document]:
        """Load DOCX using python-docx"""
        from docx import Document as DocxDocument
        
        documents = []
        doc = DocxDocument(file_path)
        
        # Extract paragraphs
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # Split by headings to create logical chunks
        current_section = ""
        current_content = []
        
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                # Save previous section
                if current_content:
                    documents.append(Document(
                        content="\n".join(current_content),
                        metadata={
                            "source": file_path,
                            "section": current_section,
                            "document_type": "docx"
                        }
                    ))
                current_section = para.text
                current_content = [para.text]
            else:
                current_content.append(para.text)
        
        # Add remaining content
        if current_content:
            documents.append(Document(
                content="\n".join(current_content),
                metadata={
                    "source": file_path,
                    "section": current_section,
                    "document_type": "docx"
                }
            ))
        
        logger.info(f"Loaded DOCX: {file_path}, {len(documents)} sections")
        return documents
    
    @staticmethod
    def _load_text(file_path: str) -> List[Document]:
        """Load plain text file"""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        return [Document(
            content=text,
            metadata={
                "source": file_path,
                "document_type": Path(file_path).suffix[1:]
            }
        )]
    
    @staticmethod
    def load_directory(dir_path: str) -> List[Document]:
        """Load all supported files from a directory"""
        documents = []
        path = Path(dir_path)
        
        for file_path in path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in DocumentLoader.SUPPORTED_EXTENSIONS:
                try:
                    docs = DocumentLoader.load_file(str(file_path))
                    documents.extend(docs)
                except Exception as e:
                    logger.error(f"Failed to load {file_path}: {e}")
        
        return documents